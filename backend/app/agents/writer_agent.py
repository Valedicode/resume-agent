"""
Writer Agent - CV and Cover Letter Tailoring Specialist

This agent receives structured data from CV and Job agents and generates:
1. Tailored CV PDFs based on job requirements
2. Cover letter PDFs aligned with the job and company

The agent follows a human-in-the-loop workflow:
- Analyzes gaps between CV and job requirements
- Proposes modifications for user approval
- Generates content after approval
- Creates professional PDF outputs

Architecture:
- Input: Pre-processed JSON data from cv_agent (ResumeInfo) and job_agent (JobRequirements, CompanyInfo)
- Processing: Gap analysis → Tailoring plan → Content generation → PDF creation
- Output: Professional PDF documents (CV and cover letter)
- Pattern: LangChain agent with specialized tools for each stage

Workflow:
1. ANALYSIS: Compare CV against job requirements (analyze_cv_job_alignment)
2. REVIEW: Present tailoring plan and wait for user approval
3. GENERATION: Create tailored CV HTML (generate_tailored_cv_html)
4. APPROVAL: Show preview and wait for user confirmation
5. PDF CREATION: Generate final CV PDF (generate_cv_pdf)
6. COVER LETTER: Generate letter content and PDF (generate_cover_letter_content, generate_cover_letter_pdf)

Key Features:
- Human-in-the-loop at every critical step
- Preserves candidate's authentic voice
- Never fabricates content - only emphasizes/refines existing material
- Professional PDF generation with table-based CV layout
- Company-aware cover letter generation
"""

from pydantic import Field
from langchain.agents import create_agent
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from pathlib import Path
import json
import html
import dotenv

# Import models from centralized schemas
from app.models.schemas import CVTailoringPlan, CoverLetterContent

dotenv.load_dotenv()

# ============================================
# WeasyPrint Import Helpers (Lazy Loading)
# ============================================

# Lazy import for WeasyPrint to avoid GTK dependency issues on Windows
# WeasyPrint will only be imported when PDF generation is actually needed
_weasyprint_available = None

def _check_weasyprint():
    """Check if WeasyPrint is available and can be imported."""
    global _weasyprint_available
    if _weasyprint_available is None:
        try:
            from weasyprint import HTML, CSS
            _weasyprint_available = True
        except (ImportError, OSError) as e:
            _weasyprint_available = False
            print(f"⚠️  WeasyPrint not available: {e}")
            print("PDF generation will not work. See installation instructions:")
            print("https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows")
    return _weasyprint_available

def _import_weasyprint():
    """Import WeasyPrint with error handling."""
    try:
        from weasyprint import HTML, CSS
        return HTML, CSS
    except (ImportError, OSError) as e:
        raise RuntimeError(
            f"WeasyPrint is not properly installed: {e}\n\n"
            "On Windows, WeasyPrint requires GTK libraries.\n"
            "Installation options:\n"
            "1. Install GTK via MSYS2: https://www.gtk.org/docs/installations/windows/\n"
            "2. Use WSL (Windows Subsystem for Linux)\n"
            "3. Use Docker\n"
            "4. Use an alternative PDF library (reportlab, fpdf)\n\n"
            "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows"
        ) from e

# ============================================
# Constants
# ============================================

# Default output directory for generated PDFs
DEFAULT_OUTPUT_DIR = Path(__file__).parent.parent.parent / "data"

# CSS Template for CV (two-column table layout)
CV_CSS_TEMPLATE = """
@page {
    size: A4;
    margin: 2cm;
}

body {
    font-family: 'Helvetica', 'Arial', sans-serif;
    font-size: 11pt;
    line-height: 1.4;
    color: #333;
}

h1 {
    font-size: 24pt;
    margin-bottom: 0.5em;
    padding-bottom: 0.3em;
    border-bottom: 2px solid #333;
}

h2 {
    font-size: 14pt;
    margin-top: 1.5em;
    margin-bottom: 0.8em;
    padding-bottom: 0.2em;
    border-bottom: 1px solid #666;
}

.contact-info {
    font-size: 10pt;
    margin-bottom: 1em;
    color: #555;
}

.section-entry {
    display: flex;
    margin-bottom: 1.2em;
    page-break-inside: avoid;
}

.entry-left {
    flex: 0 0 40%;
    font-weight: bold;
    padding-right: 1em;
}

.entry-right {
    flex: 1;
}

.entry-right .position {
    font-weight: bold;
    margin-bottom: 0.3em;
}

.entry-right .details {
    font-size: 10pt;
    color: #555;
    margin-bottom: 0.5em;
}

ul {
    margin: 0.5em 0;
    padding-left: 1.5em;
}

li {
    margin-bottom: 0.3em;
}

.skills-list {
    margin-bottom: 0.8em;
}

.skills-list strong {
    display: inline-block;
    min-width: 180px;
}
"""

# CSS Template for Cover Letter (simple single-column)
COVER_LETTER_CSS_TEMPLATE = """
@page {
    size: A4;
    margin: 2.5cm;
}

body {
    font-family: 'Times New Roman', 'Georgia', serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #000;
}

.header {
    margin-bottom: 2em;
    font-size: 11pt;
}

.date {
    margin-bottom: 2em;
}

.greeting {
    margin-bottom: 1.5em;
}

p {
    margin-bottom: 1.2em;
    text-align: justify;
}

.closing {
    margin-top: 2em;
}

.signature {
    margin-top: 3em;
}
"""

# ============================================
# Tools (Agent Capabilities)
# ============================================

@tool
def analyze_cv_job_alignment(cv_json: str, job_json: str) -> str:
    """
    Analyze how well the CV matches job requirements and create a tailoring plan.
    
    This tool performs gap analysis between the candidate's CV and the target job,
    identifying strengths to emphasize and areas where the CV can be optimized.
    
    IMPORTANT: This tool receives pre-processed JSON from cv_agent and job_agent.
    No extraction is needed - just analysis of structured data.
    
    Args:
        cv_json: ResumeInfo JSON string from cv_agent output
        job_json: JobRequirements JSON string from job_agent output
        
    Returns:
        JSON string with CVTailoringPlan containing specific, actionable suggestions
        
    Example:
        >>> plan = analyze_cv_job_alignment(cv_data, job_data)
        >>> # Returns detailed plan with matching experiences, skills, keywords, etc.
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    structured_llm = llm.with_structured_output(CVTailoringPlan)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
            You are an expert career counselor conducting gap analysis.
            Your task is to compare the candidate's CV with job requirements and identify:

            1. MATCHING EXPERIENCES: Which work experiences or roles align with job responsibilities?
            2. MATCHING SKILLS: Which technical and soft skills overlap between CV and job?
            3. RELEVANT PROJECTS: Which projects demonstrate capabilities needed for this job?
            4. KEYWORDS: What job-specific terms should appear in the tailored CV?
            5. REORDERING: Should sections be reorganized to put most relevant items first?
            6. EMPHASIS: What achievements deserve stronger highlighting?

            CRITICAL CONSTRAINTS:
            - Only work with EXISTING content from the CV
            - Never suggest fabricating experience or skills
            - Focus on emphasizing and reframing, not inventing
            - Be specific with actionable suggestions
            Provide concrete, implementable recommendations.
        """),
        ("user", """
        Candidate's CV:
        {cv}

        Job Requirements:
        {job}

        Analyze the alignment and create a detailed tailoring plan.
        """)
    ])
    
    chain = prompt | structured_llm
    result = chain.invoke({"cv": cv_json, "job": job_json})
    
    return result.model_dump_json(indent=2)

@tool
def generate_tailored_cv_html(cv_json: str, tailoring_plan_json: str) -> str:
    """
    Generate HTML content for a tailored CV based on original CV and tailoring plan.
    
    This tool creates HTML that maintains the original two-column table layout
    while incorporating the suggestions from the tailoring plan.
    
    Args:
        cv_json: Original ResumeInfo JSON from cv_agent
        tailoring_plan_json: CVTailoringPlan JSON from analyze_cv_job_alignment
        
    Returns:
        HTML string representing the tailored CV (body content only, no <html> wrapper)
        
    Note:
        The HTML uses semantic structure with classes for styling:
        - .section-entry for each CV entry
        - .entry-left for institution/company names
        - .entry-right for details and descriptions
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """
        You are generating HTML for a tailored CV.
        LAYOUT REQUIREMENTS:
        1. Maintain two-column structure:
        - Left column (40%): Institution/Company names in bold
        - Right column (60%): Roles, dates, locations, bullet points

        2. Structure:
        - <h1> for candidate name
        - <div class="contact-info"> for contact details
        - <h2> for each section (Education, Experience, Skills & Projects, etc.)
        - <div class="section-entry"> for each entry with:
            - <div class="entry-left"> (institution/company)
            - <div class="entry-right"> (details)

        3. Incorporate tailoring plan:
        - Reorder sections/items per suggestions
        - Emphasize relevant experiences and projects
        - Naturally weave in keywords from the plan
        - Highlight matching skills prominently

        4. Preserve authenticity:
        - Keep the candidate's voice and writing style
        - Don't fabricate content
        - Enhance descriptions with keywords, don't replace them

        RETURN: Only the HTML body content (no <html>, <head>, or <body> tags).
        Use semantic HTML with the specified class names for proper styling."""),
        ("user", """Original CV:
        {cv}

        Tailoring Plan:
        {plan}

        Generate the tailored CV HTML.""")
    ])
    
    chain = prompt | llm
    result = chain.invoke({"cv": cv_json, "plan": tailoring_plan_json})
    
    return result.content

@tool
def generate_cover_letter_content(cv_json: str, job_json: str, company_json: str = "") -> str:
    """
    Generate tailored cover letter content for the job application.
    
    Creates a compelling, authentic cover letter that connects the candidate's
    background to the job requirements and company culture.
    
    Args:
        cv_json: ResumeInfo JSON from cv_agent
        job_json: JobRequirements JSON from job_agent
        company_json: Optional CompanyInfo JSON from job_agent for company-specific points
        
    Returns:
        JSON string with CoverLetterContent structure (paragraph-by-paragraph)
        
    Note:
        The letter is structured in 3-4 paragraphs:
        - Opening: Express interest and how you learned about the role
        - Body 1-2: Connect relevant experiences to job requirements
        - Body 3 (optional): Address company-specific points
        - Closing: Call to action and appreciation
    """
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.4)
    structured_llm = llm.with_structured_output(CoverLetterContent)
    
    # Build context with optional company info
    company_context = ""
    if company_json and company_json.strip():
        company_context = f"\n\nCompany Information:\n{company_json}"
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are writing a professional cover letter.

        WRITING GUIDELINES:
        1. Tone: Professional yet personable, enthusiastic but not over-the-top
        2. Length: 300-400 words total (concise and impactful)
        3. Structure:
        - Opening: Hook with genuine interest in role/company
        - Body: Connect 2-3 most relevant experiences to job needs
        - Closing: Strong call to action

        4. Content principles:
        - Be SPECIFIC: Reference actual experiences and skills from CV
        - Show ALIGNMENT: Demonstrate understanding of job requirements
        - Add VALUE: Explain what you'll bring to the role
        - Be AUTHENTIC: Maintain candidate's voice

        5. If company info provided:
        - Reference company values or culture
        - Mention relevant company news or initiatives
        - Show you've done research

        AVOID:
        - Generic statements that could apply to any job
        - Simply repeating CV content
        - Explaining why YOU want the job (focus on what YOU bring)
        - Desperation or begging tone
        - Clichés and buzzwords"""),
                ("user", """Candidate's CV:
        {cv}

        Job Requirements:
        {job}
        {company}

        Write a compelling, tailored cover letter.""")
    ])
    
    chain = prompt | structured_llm
    result = chain.invoke({"cv": cv_json, "job": job_json,"company": company_context})
    
    return result.model_dump_json(indent=2)

@tool
def generate_cv_pdf(html_content: str, output_filename: str, applicant_name: str) -> str:
    """
    Generate final CV PDF from HTML content with professional styling.
    
    WARNING: Only call this tool after the user has explicitly approved the CV content!
    
    Args:
        html_content: HTML body content from generate_tailored_cv_html
        output_filename: Filename for the PDF (e.g., "kevin_ha_cv_tailored.pdf")
        applicant_name: Full name for document title metadata
        
    Returns:
        Success message with the full path to the generated PDF file
        
    Raises:
        Exception: If PDF generation fails or output directory is not writable
        
    Example:
        >>> result = generate_cv_pdf(html, "john_doe_cv.pdf", "John Doe")
        >>> print(result)
        "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
    """
    try:
        # Import WeasyPrint (lazy loading)
        HTML, CSS = _import_weasyprint()
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Build complete HTML document
        # Escape applicant_name to prevent HTML injection
        escaped_name = html.escape(applicant_name)
        full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{escaped_name} - CV</title>
        </head>
        <body>
        {html_content}
        </body>
        </html>"""
        
        # Generate PDF with styling
        HTML(string=full_html).write_pdf(
            str(output_path),
            stylesheets=[CSS(string=CV_CSS_TEMPLATE)]
        )
        
        # Return user-friendly message without exposing server path
        # Filename is included for backend detection (will be improved with tool tracking)
        return f"CV PDF generated successfully! The file '{output_filename}' is ready for download."
        
    except RuntimeError as e:
        # WeasyPrint not available
        return f"Error: {str(e)}"
    except Exception as e:
        return f"Error generating CV PDF: {str(e)}"

@tool
def generate_cover_letter_pdf(content_json: str, output_filename: str, applicant_name: str, applicant_contact: str, recipient_info: str = "Hiring Manager") -> str:
    """
    Generate cover letter PDF from structured content.
    
    Creates a simple, professional cover letter PDF with proper business letter formatting.
    
    WARNING: Only call this tool after the user has explicitly approved the cover letter content!
    
    Args:
        content_json: CoverLetterContent JSON from generate_cover_letter_content (the JSON string output from that tool)
        output_filename: Filename for the PDF (e.g., "john_doe_cover_letter.pdf")
        applicant_name: Full name from CV data - extract from cv_data['name'] or cv_data['full_name']
        applicant_contact: Contact info formatted as "email | phone" - construct from cv_data['email'] and cv_data['phone']
        recipient_info: Who to address (default: "Hiring Manager") - can extract from job_data if specified
        
    Returns:
        Success message indicating the PDF was generated (filename included for backend detection)
        
    Example:
        >>> # Get CV data from system context
        >>> name = cv_data['name']  # or cv_data['full_name']
        >>> email = cv_data.get('email', 'email@example.com')
        >>> phone = cv_data.get('phone', '')
        >>> contact = f"{email} | {phone}" if phone else email
        >>> 
        >>> result = generate_cover_letter_pdf(
        ...     content_json=cover_letter_json_output,
        ...     output_filename="john_doe_cover_letter.pdf",
        ...     applicant_name=name,
        ...     applicant_contact=contact,
        ...     recipient_info="Hiring Manager"
        ... )
    """
    try:
        # Import WeasyPrint (lazy loading)
        HTML, CSS = _import_weasyprint()
        
        # Parse content with better error handling
        try:
            content = json.loads(content_json)
        except json.JSONDecodeError as e:
            return f"Error: Invalid JSON format for cover letter content. {str(e)}"
        
        # Validate required fields
        required_fields = ['opening_paragraph', 'body_paragraph_1', 'body_paragraph_2', 'closing_paragraph']
        missing_fields = [field for field in required_fields if field not in content]
        if missing_fields:
            return f"Error: Cover letter content is missing required fields: {', '.join(missing_fields)}"
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Build paragraphs list (excluding empty optional paragraph)
        paragraphs = [
            content['opening_paragraph'],
            content['body_paragraph_1'],
            content['body_paragraph_2']
        ]
        if content.get('body_paragraph_3', '').strip():
            paragraphs.append(content['body_paragraph_3'])
        paragraphs.append(content['closing_paragraph'])
        
        # Generate paragraph HTML with proper escaping
        # Escape all paragraphs to prevent HTML injection
        paragraphs_html = '\n'.join([f'<p>{html.escape(p)}</p>' for p in paragraphs])
        
        # Get current date
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")
        
        # Escape user-provided data to prevent HTML injection
        escaped_name = html.escape(applicant_name)
        escaped_contact = html.escape(applicant_contact)
        escaped_recipient = html.escape(recipient_info)
        
        # Build complete HTML
        full_html = f"""<!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Cover Letter - {escaped_name}</title>
        </head>
        <body>
            <div class="header">
                <strong>{escaped_name}</strong><br>
                {escaped_contact}
            </div>
            
            <div class="date">
                {current_date}
            </div>
            
            <div class="greeting">
                Dear {escaped_recipient},
            </div>
            
            {paragraphs_html}
            
            <div class="closing">
                Sincerely,
            </div>
            
            <div class="signature">
                {escaped_name}
            </div>
        </body>
        </html>"""
        
        # Generate PDF with styling
        HTML(string=full_html).write_pdf(
            str(output_path),
            stylesheets=[CSS(string=COVER_LETTER_CSS_TEMPLATE)]
        )
        
        # Return user-friendly message without exposing server path
        # Filename is included for backend detection (will be improved with tool tracking)
        return f"Cover letter PDF generated successfully! The file '{output_filename}' is ready for download."
        
    except RuntimeError as e:
        # WeasyPrint not available
        return f"Error: {str(e)}"
    except Exception as e:
        return f"Error generating cover letter PDF: {str(e)}"

@tool
def generate_cv_docx(cv_json: str, output_filename: str, applicant_name: str) -> str:
    """
    Generate a Word document (.docx) with CV content in basic structure.
    
    WARNING: Only call this tool after the user has explicitly approved the CV content!
    
    This creates a simple Word document with structured content. The formatting is minimal
    so users can easily customize it in Microsoft Word or Google Docs to match their preferences.
    
    Args:
        cv_json: ResumeInfo JSON string from cv_agent (the original CV data)
        output_filename: Filename for the Word document (e.g., "kevin_ha_cv_tailored.docx")
        applicant_name: Full name for document title
    
    Returns:
        Success message with the filename for download
    
    Example:
        >>> result = generate_cv_docx(cv_data, "john_doe_cv.docx", "John Doe")
        >>> print(result)
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        # Parse CV data
        cv_data = json.loads(cv_json)
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Create document
        doc = Document()
        
        # Title (Name) - Level 1 heading
        title = doc.add_heading(cv_data.get('name', applicant_name), level=1)
        
        # Contact information
        contact_parts = []
        if cv_data.get('email'):
            contact_parts.append(cv_data['email'])
        if cv_data.get('phone'):
            contact_parts.append(cv_data['phone'])
        if cv_data.get('location'):
            contact_parts.append(cv_data['location'])
        if cv_data.get('linkedin_url'):
            contact_parts.append(f"LinkedIn: {cv_data['linkedin_url']}")
        if cv_data.get('github_url'):
            contact_parts.append(f"GitHub: {cv_data['github_url']}")
        if cv_data.get('portfolio_url'):
            contact_parts.append(f"Portfolio: {cv_data['portfolio_url']}")
        
        if contact_parts:
            contact_para = doc.add_paragraph(' | '.join(contact_parts))
            contact_para.style.font.size = Pt(10)
        
        # Education section
        if cv_data.get('education'):
            doc.add_heading('Education', level=2)
            for edu in cv_data['education']:
                edu_para = doc.add_paragraph()
                if edu.get('degree'):
                    edu_para.add_run(edu['degree']).bold = True
                if edu.get('institution'):
                    if edu.get('degree'):
                        edu_para.add_run(f" - {edu['institution']}")
                    else:
                        edu_para.add_run(edu['institution']).bold = True
                if edu.get('dates'):
                    edu_para.add_run(f" ({edu['dates']})")
                if edu.get('location'):
                    edu_para.add_run(f", {edu['location']}")
                if edu.get('gpa'):
                    edu_para.add_run(f" - GPA: {edu['gpa']}")
        
        # Experience section
        if cv_data.get('experience'):
            doc.add_heading('Experience', level=2)
            for exp in cv_data['experience']:
                # Position and company
                exp_para = doc.add_paragraph()
                if exp.get('position'):
                    exp_para.add_run(exp['position']).bold = True
                if exp.get('company'):
                    if exp.get('position'):
                        exp_para.add_run(f" at {exp['company']}")
                    else:
                        exp_para.add_run(exp['company']).bold = True
                if exp.get('dates'):
                    exp_para.add_run(f" ({exp['dates']})")
                if exp.get('location'):
                    exp_para.add_run(f" - {exp['location']}")
                
                # Responsibilities/bullet points
                if exp.get('responsibilities'):
                    for resp in exp['responsibilities']:
                        doc.add_paragraph(resp, style='List Bullet')
        
        # Skills section
        if cv_data.get('skills'):
            doc.add_heading('Skills', level=2)
            skills_text = ', '.join(cv_data['skills'])
            doc.add_paragraph(skills_text)
        
        # Projects section
        if cv_data.get('projects'):
            doc.add_heading('Projects', level=2)
            for project in cv_data['projects']:
                proj_para = doc.add_paragraph()
                if project.get('name'):
                    proj_para.add_run(project['name']).bold = True
                if project.get('description'):
                    if project.get('name'):
                        proj_para.add_run(f": {project['description']}")
                    else:
                        proj_para.add_run(project['description'])
                if project.get('technologies'):
                    tech_text = ', '.join(project['technologies']) if isinstance(project['technologies'], list) else project['technologies']
                    doc.add_paragraph(f"Technologies: {tech_text}", style='List Bullet')
        
        # Leadership & Activities section
        if cv_data.get('leadership_activities'):
            doc.add_heading('Leadership & Activities', level=2)
            for activity in cv_data['leadership_activities']:
                act_para = doc.add_paragraph()
                if activity.get('role'):
                    act_para.add_run(activity['role']).bold = True
                if activity.get('organization'):
                    if activity.get('role'):
                        act_para.add_run(f" - {activity['organization']}")
                    else:
                        act_para.add_run(activity['organization']).bold = True
                if activity.get('dates'):
                    act_para.add_run(f" ({activity['dates']})")
                if activity.get('description'):
                    doc.add_paragraph(activity['description'], style='List Bullet')
        
        # Save document
        doc.save(str(output_path))
        
        return f"CV Word document generated successfully! The file '{output_filename}' is ready for download. You can customize the formatting in Microsoft Word or Google Docs."
        
    except ImportError:
        return "Error: python-docx library is not installed. Please install it to generate Word documents."
    except json.JSONDecodeError as e:
        return f"Error: Invalid CV JSON data: {str(e)}"
    except Exception as e:
        return f"Error generating CV Word document: {str(e)}"

@tool
def generate_cover_letter_docx(content_json: str, output_filename: str, applicant_name: str, applicant_contact: str, recipient_info: str = "Hiring Manager") -> str:
    """
    Generate a Word document (.docx) with cover letter content in basic structure.
    
    WARNING: Only call this tool after the user has explicitly approved the cover letter content!
    
    This creates a simple Word document with structured cover letter content. The formatting is minimal
    so users can easily customize it in Microsoft Word or Google Docs.
    
    Args:
        content_json: CoverLetterContent JSON from generate_cover_letter_content (the JSON string output)
        output_filename: Filename for the Word document (e.g., "john_doe_cover_letter.docx")
        applicant_name: Full name from CV data
        applicant_contact: Contact info formatted as "email | phone"
        recipient_info: Who to address (default: "Hiring Manager")
    
    Returns:
        Success message with the filename for download
    
    Example:
        >>> result = generate_cover_letter_docx(
        ...     content_json=cover_letter_json_output,
        ...     output_filename="john_doe_cover_letter.docx",
        ...     applicant_name="John Doe",
        ...     applicant_contact="john@email.com | (123) 456-7890)",
        ...     recipient_info="Hiring Manager"
        ... )
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from datetime import datetime
        
        # Parse cover letter content
        content_data = json.loads(content_json)
        
        # Ensure output directory exists
        output_dir = DEFAULT_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Construct full output path
        output_path = output_dir / output_filename
        
        # Create document
        doc = Document()
        
        # Header with name and contact
        header_para = doc.add_paragraph()
        header_para.add_run(applicant_name).bold = True
        header_para.add_run(f"\n{applicant_contact}")
        
        # Date
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        current_date = datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph(current_date)
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Greeting
        greeting_para = doc.add_paragraph()
        greeting_para.add_run(f"Dear {recipient_info},")
        
        # Body paragraphs
        if content_data.get('paragraphs'):
            for para_text in content_data['paragraphs']:
                doc.add_paragraph(para_text)
        
        # Closing
        closing_para = doc.add_paragraph("Sincerely,")
        closing_para.space_after = Pt(12)
        
        # Signature
        signature_para = doc.add_paragraph(applicant_name)
        signature_para.space_before = Pt(24)
        
        # Save document
        doc.save(str(output_path))
        
        return f"Cover letter Word document generated successfully! The file '{output_filename}' is ready for download. You can customize the formatting in Microsoft Word or Google Docs."
        
    except ImportError:
        return "Error: python-docx library is not installed. Please install it to generate Word documents."
    except json.JSONDecodeError as e:
        return f"Error: Invalid cover letter JSON data: {str(e)}"
    except Exception as e:
        return f"Error generating cover letter Word document: {str(e)}"

# ============================================
# Agent Configuration
# ============================================

# Initialize LLM for the agent
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.2)

# Comprehensive system prompt defining agent behavior
system_prompt = """You are a professional CV and cover letter writer specializing in tailoring application materials to specific job opportunities.

IMPORTANT CONTEXT:
- You receive PRE-PROCESSED, STRUCTURED DATA from upstream agents:
  * CV data (ResumeInfo JSON) - extracted by cv_agent
  * Job requirements (JobRequirements JSON) - extracted by job_agent  
  * Company info (CompanyInfo JSON, optional) - researched by job_agent
- Your role is to TRANSFORM this data into tailored, professional documents
- You DO NOT extract data from raw files - that's already done
- The complete CV and job data are provided in a system message at the start of our conversation
- This data persists throughout our entire conversation - always use the FULL JSON data from the system context when calling tools
- Do NOT use truncated or partial data from earlier messages - always reference the complete data from the system context

YOUR WORKFLOW (FOLLOW THIS SEQUENCE):

1. ANALYSIS PHASE:
   - When user provides CV and job data, call analyze_cv_job_alignment
   - This generates a tailoring plan showing how to optimize the CV

2. REVIEW PHASE:
   - PRESENT the tailoring plan to the user in clear, readable format
   - IMPORTANT: Avoid complex markdown structure (no ### or deeper headers, no - lists, no code blocks)
   - You may use simple formatting: ## for section headers, **bold** for emphasis, and *italic* for subtle emphasis
   - Use simple line breaks, indentation, and numbered/bullet points written as plain text
   - Explain the key recommendations
   - WAIT for user feedback before proceeding
   - User might say:
     * "looks good" or "approved" → proceed to generation
     * "emphasize X more" → note the adjustment
     * "don't mention Y" → note the constraint
     * "redo the analysis" → call the tool again with updated context
   - After refining or discussing each section (Education, Experience, Skills, Projects, etc.), ask if they want to generate the document for that section
   - You can refine multiple sections before generating, but always offer generation after showing each refined section

3. CV GENERATION PHASE:
   - After approval, call generate_tailored_cv_html
   - SHOW a preview/summary of the generated HTML to the user
   - After showing each section or the complete CV, ALWAYS ask: "Ready to generate the document? I can create a PDF, Word document (.docx), or both formats. Which would you prefer?"
   - If user requests "both", "pdf and word", "both formats", or similar, generate BOTH formats
   - If user requests "pdf" or "PDF", generate only PDF
   - If user requests "word", "docx", or "Word", generate only Word document
   - If user doesn't specify, ask for clarification
   - ONLY call generate_cv_pdf and/or generate_cv_docx after EXPLICIT approval
   - Use appropriate filename (e.g., "firstname_lastname_cv_tailored.pdf" or ".docx")
   - When generating both formats, call BOTH tools and include BOTH success messages in your response
   - Note: Word documents have minimal formatting so users can easily customize them

4. COVER LETTER PHASE (if requested):
   - Call generate_cover_letter_content with CV, job, and company data
   - SHOW the cover letter content to the user
   - Ask for feedback or approval
   - After showing the cover letter, ALWAYS ask: "Ready to generate the document? I can create a PDF, Word document (.docx), or both formats. Which would you prefer?"
   - If user requests "both", "pdf and word", "both formats", or similar, generate BOTH formats
   - If user requests "pdf" or "PDF", generate only PDF
   - If user requests "word", "docx", or "Word", generate only Word document
   - If user doesn't specify, ask for clarification
   - ONLY call generate_cover_letter_pdf and/or generate_cover_letter_docx after EXPLICIT approval
   - When calling generate_cover_letter_pdf or generate_cover_letter_docx, you MUST provide:
     * content_json: The output from generate_cover_letter_content
     * output_filename: e.g., "firstname_lastname_cover_letter.pdf" or ".docx"
     * applicant_name: Extract from CV data (e.g., cv_data['name'])
     * applicant_contact: Format as "email | phone" from CV data (e.g., "john@email.com | (123) 456-7890")
     * recipient_info: Extract from job data or use "Hiring Manager"
   - When generating both formats, call BOTH tools and include BOTH success messages in your response

CRITICAL PRINCIPLES:
ALWAYS show changes before generating PDFs
WAIT for explicit user approval at each checkpoint
Preserve the candidate's authentic voice and style
Only emphasize/refine EXISTING content - never fabricate
Be transparent about what you're changing and why
Respect user constraints and preferences
If uncertain, ask clarifying questions

NEVER generate PDFs without approval
NEVER fabricate experience, skills, or achievements
NEVER proceed to next phase without user confirmation
NEVER ignore user feedback or constraints

INTERACTION STYLE:
- Professional yet friendly
- Clear and concise in explanations
- Proactive in asking for clarification
- Transparent about your process
- Collaborative, not autonomous
- Avoid complex markdown structure (no ### or deeper headers, no - lists, no code blocks)
- You may use simple formatting: ## for section headers, **bold** for emphasis, and *italic* for subtle emphasis
- Format responses with simple line breaks and plain text structure

**CRITICAL - File Downloads:**
- When you call generate_cv_pdf, generate_cv_docx, generate_cover_letter_pdf, or generate_cover_letter_docx tools, they return a success message
- You MUST include this EXACT success message in your response to the user
- DO NOT paraphrase or reword the tool's success message - copy it verbatim
- The success message contains the filename which is needed for the download system to work
- DO NOT create markdown links like [Download](sandbox:/filename.pdf) or [Download](file://filename.pdf)
- DO NOT create any file://, sandbox:/, or other file links - they don't work
- The user interface will AUTOMATICALLY display download buttons when it detects the tool's success messages
- If you generate BOTH formats (PDF and Word), include BOTH success messages and the UI will show TWO download buttons
- You can add additional text before or after the tool messages, but each tool's exact message MUST be included
- For Word documents, mention that formatting is minimal so users can customize it easily

Examples:
- Single format: Tool returns "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
  Your response: "Great! CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
- Both formats: 
  Tool 1 returns: "CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download."
  Tool 2 returns: "CV Word document generated successfully! The file 'john_doe_cv.docx' is ready for download."
  Your response MUST include BOTH messages: "Perfect! I've generated both formats for you. CV PDF generated successfully! The file 'john_doe_cv.pdf' is ready for download. CV Word document generated successfully! The file 'john_doe_cv.docx' is ready for download. You can customize the Word document formatting in Microsoft Word or Google Docs."

Remember: You are an assistant helping the user create their best application materials. 
The user is the expert on their own experience - you're the expert on presentation."""

# Create the Writer Agent
# This agent uses LangChain's create_agent which provides:
# - Automatic tool selection based on user input
# - Conversation memory across turns
# - Structured reasoning about when to use each tool
agent = create_agent(
    model=llm,
    tools=[
        analyze_cv_job_alignment,
        generate_tailored_cv_html,
        generate_cover_letter_content,
        generate_cv_pdf,
        generate_cover_letter_pdf,
        generate_cv_docx,
        generate_cover_letter_docx
    ],
    system_prompt=system_prompt
)

# ============================================
# Interactive Runner (for testing)
# ============================================

def run_interactive_writer(cv_json: str = None, job_json: str = None, company_json: str = None):
    """
    Run writer agent interactively in terminal for testing.
    
    This function provides a CLI interface to test the writer agent with
    mock or real data from cv_agent and job_agent outputs.
    
    Args:
        cv_json: Optional ResumeInfo JSON string (if None, prompts user)
        job_json: Optional JobRequirements JSON string (if None, prompts user)
        company_json: Optional CompanyInfo JSON string (can be None)
        
    Returns:
        Dict with conversation history and results
        
    Example:
        >>> # With data
        >>> run_interactive_writer(cv_data, job_data)
        
        >>> # Interactive input
        >>> run_interactive_writer()
    """
    print("=" * 70)
    print("WRITER AGENT - Interactive Testing Mode")
    print("=" * 70)
    print("\nThis agent helps you tailor CVs and cover letters to job opportunities.")
    print("It works with structured JSON data from cv_agent and job_agent.\n")
    
    # Get data if not provided
    if not cv_json:
        print("Enter path to CV JSON file (or 'mock' for test data):")
        cv_input = input("> ").strip()
        if cv_input.lower() == 'mock':
            cv_json = json.dumps({
                "name": "Test Candidate",
                "email": "test@example.com",
                "phone": "+1234567890",
                "skills": ["Python", "JavaScript", "React"],
                "education": ["B.Sc. Computer Science"],
                "experience": ["Software Developer at Tech Corp"]
            })
        else:
            with open(cv_input, 'r') as f:
                cv_json = f.read()
    
    if not job_json:
        print("\nEnter path to Job JSON file (or 'mock' for test data):")
        job_input = input("> ").strip()
        if job_input.lower() == 'mock':
            job_json = json.dumps({
                "job_title": "Senior Software Engineer",
                "job_level": "Senior",
                "required_skills": ["Python", "React", "AWS"],
                "responsibilities": ["Build scalable applications"],
                "qualifications": ["Bachelor's degree in CS"]
            })
        else:
            with open(job_input, 'r') as f:
                job_json = f.read()
    
    # Initialize conversation
    initial_message = f"""I have CV and job data ready for tailoring.

CV Data:
{cv_json[:200]}...

Job Data:
{job_json[:200]}...

Please analyze the alignment and help me create a tailored CV."""
    
    messages = [{"role": "user", "content": initial_message}]
    
    # Conversation loop
    max_turns = 20
    turn = 0
    
    print("\n" + "=" * 70)
    print("Starting Conversation (type 'quit' to exit)")
    print("=" * 70)
    
    while turn < max_turns:
        turn += 1
        print(f"\n{'='*70}")
        print(f"Turn {turn}")
        print(f"{'='*70}\n")
        
        # Agent's turn
        print("🤖 Agent is processing...\n")
        try:
            result = agent.invoke({"messages": messages})
            agent_response = result["messages"][-1].content
            
            # Add to history
            messages.append({"role": "assistant", "content": agent_response})
            
            # Display response
            print(f"Agent:\n{agent_response}\n")
            
        except Exception as e:
            print(f"❌ Error: {e}")
            import traceback
            traceback.print_exc()
            break
        
        # User's turn
        print(f"{'-'*70}")
        user_input = input("You: ").strip()
        
        # Check for exit
        if user_input.lower() in ['quit', 'exit', 'q']:
            print("\n👋 Exiting writer agent.")
            break
        
        # Add user response
        messages.append({"role": "user", "content": user_input})
    
    if turn >= max_turns:
        print(f"\n⚠️  Reached maximum turns ({max_turns}).")
    
    return {
        "messages": messages,
        "turns": turn
    }

# ============================================
# CLI Entry Point
# ============================================

if __name__ == "__main__":
    import sys
    
    print("\n" + "="*70)
    print("WRITER AGENT - Command Line Interface")
    print("="*70 + "\n")
    
    # Check for command line arguments
    if len(sys.argv) > 2:
        cv_path = sys.argv[1]
        job_path = sys.argv[2]
        company_path = sys.argv[3] if len(sys.argv) > 3 else None
        
        print(f"Loading CV data from: {cv_path}")
        print(f"Loading Job data from: {job_path}")
        if company_path:
            print(f"Loading Company data from: {company_path}")
        
        try:
            with open(cv_path, 'r') as f:
                cv_json = f.read()
            with open(job_path, 'r') as f:
                job_json = f.read()
            company_json = None
            if company_path:
                with open(company_path, 'r') as f:
                    company_json = f.read()
            
            run_interactive_writer(cv_json, job_json, company_json)
            
        except FileNotFoundError as e:
            print(f"\n❌ Error: File not found - {e}")
            sys.exit(1)
        except Exception as e:
            print(f"\n❌ Error: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)
    else:
        print("Usage:")
        print("  python writer_agent.py <cv_json_path> <job_json_path> [company_json_path]")
        print("\nOr run without arguments for interactive mode:")
        print("  python writer_agent.py\n")
        
        proceed = input("Run in interactive mode? (y/n): ").strip().lower()
        if proceed == 'y':
            try:
                run_interactive_writer()
            except KeyboardInterrupt:
                print("\n\n👋 Interrupted by user.")
            except Exception as e:
                print(f"\n❌ Error: {e}")
                import traceback
                traceback.print_exc()
        else:
            print("Exiting.")